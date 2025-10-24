"""Generate DOCX files for memo drafts."""
from __future__ import annotations

from pathlib import Path
from typing import Any, Dict

from docx import Document


class DocxBuilder:
    def __init__(self, base_dir: Path | None = None) -> None:
        self.base_dir = base_dir or Path(".docs")
        self.base_dir.mkdir(parents=True, exist_ok=True)

    def build(self, deal_id: str, memo: Dict[str, Any]) -> Path:
        document = Document()
        document.add_heading(f"Investment Memo Draft - {deal_id}", level=0)

        for section, payload in memo.items():
            document.add_heading(section.replace("_", " ").title(), level=1)
            self._render_section(document, payload)

        output_path = self.base_dir / f"{deal_id}_memo.docx"
        document.save(output_path)
        return output_path

    def _render_section(self, document: Document, payload: Any, indent: int = 0) -> None:
        if isinstance(payload, dict):
            for key, value in payload.items():
                document.add_paragraph(f"{key.replace('_', ' ').title()}:")
                self._render_section(document, value, indent + 1)
        elif isinstance(payload, list):
            for item in payload:
                self._render_section(document, item, indent)
        else:
            document.add_paragraph(str(payload))

